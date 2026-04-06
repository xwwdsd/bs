#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from __future__ import annotations

import json
import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph
from PIL import Image


ROOT = Path(r"E:\bs")
SOURCE_DOC = ROOT / "_doc_extract" / "初稿.docx"
TARGET_DOC = ROOT / "_doc_extract" / "初稿-插图完成版-v3.docx"
MANIFEST = ROOT / "_doc_extract" / "generated_figures" / "manifest.json"


def insert_paragraph_after(paragraph: Paragraph, text: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    return new_para


def clear_paragraph(paragraph: Paragraph) -> None:
    p = paragraph._element
    for child in list(p):
        p.remove(child)


def find_paragraph(doc: Document, needle: str) -> Paragraph:
    for paragraph in doc.paragraphs:
        if needle in paragraph.text:
            return paragraph
    raise ValueError(f"未找到段落: {needle}")


def figure_width_inches(section, image_path: Path) -> float:
    text_width = float(section.page_width - section.left_margin - section.right_margin) / 914400.0
    text_height = float(section.page_height - section.top_margin - section.bottom_margin) / 914400.0
    with Image.open(image_path) as im:
        width_px, height_px = im.size
    ratio = width_px / height_px

    if ratio >= 2.0:
        width_in = text_width * 0.98
    elif ratio >= 1.2:
        width_in = text_width * 0.90
    elif ratio >= 0.8:
        width_in = text_width * 0.84
    else:
        width_in = text_width * 0.72

    max_height_in = text_height * 0.70
    predicted_height_in = width_in / ratio
    if predicted_height_in > max_height_in:
        width_in = max_height_in * ratio

    return width_in


def add_picture_block(paragraph: Paragraph, image_path: Path, caption: str, width_in: float, reuse_paragraph: bool = False) -> Paragraph:
    if reuse_paragraph:
        pic_para = paragraph
        clear_paragraph(pic_para)
    else:
        pic_para = insert_paragraph_after(paragraph)
    pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = pic_para.add_run()
    run.add_picture(str(image_path), width=Inches(width_in))

    caption_para = insert_paragraph_after(pic_para, caption)
    caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in caption_para.runs:
        run.font.name = "Times New Roman"
        if run._element.rPr is None:
            run._element.get_or_add_rPr()
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        run.font.size = Pt(10.5)

    spacer = insert_paragraph_after(caption_para)
    return spacer
def main() -> int:
    shutil.copyfile(SOURCE_DOC, TARGET_DOC)
    doc = Document(TARGET_DOC)
    manifest = json.loads(MANIFEST.read_text(encoding="utf-8"))
    figure_map = {item["number"]: item for item in manifest["figures"]}
    figure_map["4-1"] = manifest["aliases"]["4-1"]

    section = doc.sections[0]

    def width(num: str) -> float:
        return figure_width_inches(section, Path(figure_map[num]["file"]))

    # After headings
    anchor = find_paragraph(doc, "2.1.1 买家用户需求分析")
    add_picture_block(anchor, Path(figure_map["2-1"]["file"]), figure_map["2-1"]["caption"], width("2-1"))

    anchor = find_paragraph(doc, "2.1.3 管理员需求分析")
    add_picture_block(anchor, Path(figure_map["2-2"]["file"]), figure_map["2-2"]["caption"], width("2-2"))

    anchor = find_paragraph(doc, "3.3 买家用户功能设计")
    add_picture_block(anchor, Path(figure_map["3-3"]["file"]), figure_map["3-3"]["caption"], width("3-3"))

    anchor = find_paragraph(doc, "3.4.1 Steam 绑定与库存同步设计")
    add_picture_block(anchor, Path(figure_map["3-4"]["file"]), figure_map["3-4"]["caption"], width("3-4"))

    anchor = find_paragraph(doc, "3.3.2 饰品详情与订单创建设计")
    add_picture_block(anchor, Path(figure_map["3-5"]["file"]), figure_map["3-5"]["caption"], width("3-5"))

    anchor = find_paragraph(doc, "3.4.2 出售订单与发货流程设计")
    add_picture_block(anchor, Path(figure_map["3-6"]["file"]), figure_map["3-6"]["caption"], width("3-6"))

    anchor = find_paragraph(doc, "3.4.3 收益管理设计")
    add_picture_block(anchor, Path(figure_map["3-7"]["file"]), figure_map["3-7"]["caption"], width("3-7"))

    anchor = find_paragraph(doc, "3.5.2 内容审核与统计设计")
    add_picture_block(anchor, Path(figure_map["3-8"]["file"]), figure_map["3-8"]["caption"], width("3-8"))

    anchor = find_paragraph(doc, "3.5.1 平台运营管理设计")
    add_picture_block(anchor, Path(figure_map["3-9"]["file"]), figure_map["3-9"]["caption"], width("3-9"))

    anchor = find_paragraph(doc, "3.6.2 用户实体设计")
    add_picture_block(anchor, Path(figure_map["3-11"]["file"]), figure_map["3-11"]["caption"], width("3-11"))

    anchor = find_paragraph(doc, "3.6.3 饰品实体设计")
    add_picture_block(anchor, Path(figure_map["3-12"]["file"]), figure_map["3-12"]["caption"], width("3-12"))

    anchor = find_paragraph(doc, "3.6.5 出售订单与交易订单实体设计")
    add_picture_block(anchor, Path(figure_map["3-13"]["file"]), figure_map["3-13"]["caption"], width("3-13"))

    anchor = find_paragraph(doc, "3.6.6 钱包与辅助实体设计")
    spacer = add_picture_block(anchor, Path(figure_map["3-14"]["file"]), figure_map["3-14"]["caption"], width("3-14"))
    add_picture_block(spacer, Path(figure_map["3-15"]["file"]), figure_map["3-15"]["caption"], width("3-15"))

    # Replace placeholders
    placeholder = find_paragraph(doc, "图3-1 系统总体架构图")
    add_picture_block(placeholder, Path(figure_map["3-1"]["file"]), figure_map["3-1"]["caption"], width("3-1"), reuse_paragraph=True)

    placeholder = find_paragraph(doc, "图3-2 系统功能模块图")
    add_picture_block(placeholder, Path(figure_map["3-2"]["file"]), figure_map["3-2"]["caption"], width("3-2"), reuse_paragraph=True)

    placeholder = find_paragraph(doc, "图3-3 数据库 E-R 关系图")
    add_picture_block(placeholder, Path(figure_map["3-10"]["file"]), figure_map["3-10"]["caption"], width("3-10"), reuse_paragraph=True)

    placeholder = find_paragraph(doc, "图4-1 交易订单处理流程图")
    add_picture_block(placeholder, Path(figure_map["4-1"]["file"]), figure_map["4-1"]["caption"], width("4-1"), reuse_paragraph=True)

    doc.save(TARGET_DOC)
    print(TARGET_DOC)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
